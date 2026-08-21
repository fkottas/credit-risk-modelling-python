"""Build the review DOCX from Markdown.

Preset: compact_reference_guide. Header pattern: editorial_cover.
Named overrides: cover, part dividers, code blocks, reference text, wide tables.
"""

from __future__ import annotations

import argparse
import copy
import json
import re
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "book" / "full_manuscript"
STRUCTURE = ROOT / "book" / "structure.json"
GUIDED_LABS = ROOT / "book" / "guided_labs"
FIGURES = ROOT / "book" / "figures"
PAGE_MAP = ROOT / "book" / "page_map.json"
BLUE, DARK_BLUE, INK = "2E74B5", "1F4D78", "203748"
SUBTITLE, GOLD, MUTED = "2B5163", "8B6F28", "667788"
TABLE_FILL, LIGHT_FILL, CALLOUT_FILL = "E8EEF5", "F2F4F7", "F4F6F9"
PAGE_WIDTH_DXA = 9360
MATH_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
EQUATION_CATALOG: PandocEquationCatalog | None = None
ROMAN_MATH_TOKENS = (
    "CountContracts",
    "CitationPrecision",
    "SupportRate",
    "ChangeHash",
    "UnexpectedLoss",
    "Reconciliation",
    "BaseScore",
    "BaseOdds",
    "CumPD",
    "maxDPD",
    "LossRate",
    "Entropy",
    "Brier",
    "logit",
    "Gini",
    "Gain",
    "RAROC",
    "Release",
    "Allowed",
    "Policy",
    "Security",
    "Approval",
    "Tests",
    "Trigger",
    "Finding",
    "Downturn",
    "Revenue",
    "Economic",
    "Capital",
    "Opening",
    "Closing",
    "Offset",
    "Factor",
    "Deploy",
    "Score",
    "Stage",
    "Floor",
    "SHA256",
    "CECL",
    "ECL",
    "MPD",
    "mPD",
    "LGD",
    "EAD",
    "SICR",
    "WOE",
    "PDO",
    "RWA",
    "CCF",
    "EIR",
    "CVA",
    "BM25",
    "MoC",
    "LRA",
    "OOT",
    "AUC",
    "PSI",
    "SSE",
    "Beta",
    "criterion",
    "evidence",
    "severity",
    "watchlist",
    "thresholds",
    "mandatory",
    "violations",
    "historical",
    "supported",
    "critical",
    "all",
    "previous",
    "approval",
    "exposure",
    "writeoffs",
    "transfers",
    "threshold",
    "default",
    "claims",
    "record",
    "change",
    "metric",
    "config",
    "action",
    "scope",
    "owner",
    "losses",
    "credit",
    "cited",
    "charge",
    "ratio",
    "hash",
    "role",
    "time",
    "date",
    "data",
    "code",
    "Cost",
    "pass",
    "due",
    "FX",
    "UL",
    "EL",
    "PD",
    "IV",
    "DF",
    "EE",
    "KS",
)
ROMAN_MATH_PATTERN = re.compile(
    rf"(?<![A-Za-z\\])({'|'.join(map(re.escape, ROMAN_MATH_TOKENS))})(?![A-Za-z])"
)
TEXT_MATH_PATTERN = re.compile(r"\\text\{[^{}]*\}")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:tcW"))
    if node is None:
        node = OxmlElement("w:tcW")
        tc_pr.append(node)
    node.set(qn("w:w"), str(width))
    node.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, color: str = "C8D4DF", size: int = 4) -> None:
    """Apply a restrained single grid that survives Word and PDF rendering."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_picture_alt_text(inline_shape, description: str) -> None:
    """Attach meaningful alternative text to an inline Word image."""
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", description)
    doc_pr.set("descr", description)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), DARK_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    properties.extend([color, underline, size])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([properties, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


INLINE = re.compile(r"(\$[^$\n]+\$|\*\*.*?\*\*|\*[^*]+\*|`.*?`|https?://[^\s)]+[\w/#])")


def add_inline(paragraph, text: str, *, size: float | None = None) -> None:
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position : match.start()]), size=size)
        token = match.group(0)
        if token.startswith("$"):
            if EQUATION_CATALOG is None:
                raise RuntimeError("The native Word equation catalogue has not been initialised")
            EQUATION_CATALOG.append(paragraph, token[1:-1])
        elif token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=size, bold=True)
        elif token.startswith("*"):
            set_run_font(paragraph.add_run(token[1:-1]), size=size, italic=True)
        elif token.startswith("`"):
            set_run_font(
                paragraph.add_run(token[1:-1]), name="Consolas", size=size or 9.5, color=DARK_BLUE
            )
        else:
            add_hyperlink(paragraph, token, token)
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]), size=size)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True
    # Body paragraphs must be allowed to flow across pages. Keeping every
    # paragraph together created large artificial gaps throughout the book.
    normal.paragraph_format.keep_together = False
    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name, style.font.size, style.font.bold = "Calibri", Pt(size), True
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before, style.paragraph_format.space_after = (
            Pt(before),
            Pt(after),
        )
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
    # Parts, rather than all 72 short teaching chapters, control page starts.
    doc.styles["Heading 1"].paragraph_format.page_break_before = False
    for name in ("Header", "Footer"):
        style = doc.styles[name]
        style.font.name, style.font.size = "Calibri", Pt(8.5)
        style.font.color.rgb = RGBColor.from_string(MUTED)
        style.paragraph_format.space_after = Pt(0)


def add_custom_numbering(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abs_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id, num_id = max(abs_ids, default=-1) + 1, max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (
        ("start", "1"),
        ("numFmt", "bullet" if bullet else "decimal"),
        ("lvlText", "•" if bullet else "%1."),
        ("lvlJc", "left"),
    ):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        level.append(node)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    for tag, value in (("ilvl", "0"), ("numId", str(num_id))):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        num_pr.append(node)
    p_pr.append(num_pr)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    for tag, value in (("fldChar", "begin"), ("instrText", " PAGE "), ("fldChar", "separate")):
        node = OxmlElement(f"w:{tag}")
        if tag == "fldChar":
            node.set(qn("w:fldCharType"), value)
        else:
            node.set(qn("xml:space"), "preserve")
            node.text = value
        run._r.append(node)
    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = (
        Inches(1)
    )
    section.header_distance = section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    paragraph = section.header.paragraphs[0]
    run = paragraph.add_run("INTELLIGENT CREDIT RISK MODELING WITH PYTHON  •  TEACHING EDITION")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    add_page_field(section.footer.paragraphs[0])
    section.first_page_header.paragraphs[0].clear()
    section.first_page_footer.paragraphs[0].clear()


def add_cover(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    kicker = doc.add_paragraph()
    kicker.alignment, kicker.paragraph_format.space_after = WD_ALIGN_PARAGRAPH.CENTER, Pt(3)
    set_run_font(
        kicker.add_run("AN ANALYTICAL AND APPLIED TEACHING EDITION"),
        size=8.5,
        color=GOLD,
        bold=True,
    )
    title = doc.add_paragraph()
    title.alignment, title.paragraph_format.space_after = WD_ALIGN_PARAGRAPH.CENTER, Pt(1)
    set_run_font(title.add_run("Intelligent Credit Risk"), size=23, color=INK, bold=True)
    second_title = doc.add_paragraph()
    second_title.alignment, second_title.paragraph_format.space_after = (
        WD_ALIGN_PARAGRAPH.CENTER,
        Pt(2),
    )
    set_run_font(second_title.add_run("Modeling with Python"), size=23, color=INK, bold=True)
    author = doc.add_paragraph()
    author.alignment, author.paragraph_format.space_after = WD_ALIGN_PARAGRAPH.CENTER, Pt(2)
    set_run_font(
        author.add_run("Dr. Ferdinantos Kottas  •  August 2026"),
        size=10.5,
        color=SUBTITLE,
        bold=True,
    )
    cover_art = FIGURES / "cover-intelligent-credit-risk.png"
    if cover_art.exists():
        picture = doc.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.paragraph_format.space_before = Pt(1)
        picture.paragraph_format.space_after = Pt(0)
        shape = picture.add_run().add_picture(str(cover_art), width=Inches(5.20))
        set_picture_alt_text(
            shape,
            "Book-cover illustration of a governed end-to-end intelligent credit-risk system.",
        )
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    title = doc.add_paragraph("Contents", style="Heading 1")
    title.paragraph_format.page_break_before = False
    add_bookmark(title, "book_contents", 10000)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    set_run_font(
        note.add_run("Page numbers and chapter links are generated from the document headings."),
        size=9.5,
        color=MUTED,
        italic=True,
    )
    structure = json.loads(STRUCTURE.read_text(encoding="utf-8"))
    page_map = json.loads(PAGE_MAP.read_text(encoding="utf-8")) if PAGE_MAP.exists() else {}
    roman = ("I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII")
    entries: list[tuple[str, str, str, str | None, bool]] = []
    for index, part in enumerate(structure["parts"]):
        first_chapter = part["chapters"][0]["number"]
        entries.append(
            (
                f"PART {roman[index]}",
                part["title"],
                str(page_map.get(f"Chapter {first_chapter}", "")),
                f"chapter_{first_chapter:03d}",
                True,
            )
        )
        for chapter in part["chapters"]:
            number = chapter["number"]
            entries.append(
                (
                    f"Chapter {number}",
                    chapter["title"],
                    str(page_map.get(f"Chapter {number}", "")),
                    f"chapter_{number:03d}",
                    False,
                )
            )
    for label, title_text, heading in (
        (
            "Appendices",
            "APIs, policies, templates, legal data catalogue and references",
            "Appendices",
        ),
        (
            "Casebook",
            "Seventy-two worked assignments and evidence requirements",
            "Practice Casebook — Seventy-Two Worked Assignments",
        ),
        (
            "Workbook",
            "Twelve end-to-end Python implementation workshops",
            "Technical Workbook — End-to-End Python Patterns",
        ),
        (
            "Numerical",
            "Twelve hand-auditable calculation examples",
            "Numerical Examples — Calculation, Interpretation, and Audit",
        ),
        (
            "Policies",
            "Sixteen data, model, accounting, capital and AI policies",
            "Credit Risk Policy Playbook",
        ),
        (
            "Review",
            "Seventy-two viva questions with instructor notes",
            "Review and Viva Questions with Instructor Notes",
        ),
        (
            "Glossary",
            "Technical terms and control-language reference",
            "Technical and Governance Glossary",
        ),
    ):
        anchor = re.sub(r"[^a-z0-9]+", "_", heading.lower()).strip("_")
        entries.append(
            (label, title_text, str(page_map.get(heading, "")), f"section_{anchor}", False)
        )

    table = doc.add_table(rows=len(entries), cols=3)
    table_geometry(table, [1450, 7150, 760], indent_dxa=0)
    set_table_borders(table, color="D7E0E8", size=2)
    for row, (label, item, page, anchor, is_part) in zip(table.rows, entries, strict=False):
        for column, cell in enumerate(row.cells):
            set_cell_width(cell, (1450, 7150, 760)[column])
            set_cell_margins(cell, top=20, start=35, bottom=20, end=45)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
        if anchor:
            add_internal_hyperlink(row.cells[0].paragraphs[0], label, anchor)
            add_internal_hyperlink(row.cells[1].paragraphs[0], item, anchor)
            if page:
                add_internal_hyperlink(row.cells[2].paragraphs[0], page, anchor)
        else:
            set_run_font(
                row.cells[0].paragraphs[0].add_run(label),
                size=8.5,
                color=GOLD if is_part else DARK_BLUE,
                bold=True,
            )
            set_run_font(row.cells[1].paragraphs[0].add_run(item), size=8.5)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if not anchor or not page:
            set_run_font(row.cells[2].paragraphs[0].add_run(page), size=8.5, color=MUTED)
        if is_part:
            for cell in row.cells:
                set_cell_shading(cell, CALLOUT_FILL)
    doc.add_page_break()


def set_paragraph_bottom_border(paragraph, *, color: str = "C8D4DF", size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)


def add_part_banner(doc: Document, label: str, title: str, *, first_part: bool) -> None:
    """Start a part compactly on the same page as its first chapter."""

    part = doc.add_paragraph()
    part.paragraph_format.page_break_before = False
    part.paragraph_format.space_before = Pt(0 if first_part else 16)
    part.paragraph_format.space_after = Pt(2)
    part.paragraph_format.keep_with_next = True
    set_run_font(part.add_run(label.upper()), size=9.0, color=GOLD, bold=True)
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(10)
    heading.paragraph_format.keep_with_next = True
    set_run_font(heading.add_run(title), size=15.5, color=INK, bold=True)
    set_paragraph_bottom_border(heading, color="B9C9D8", size=8)


def table_geometry(table, widths: list[int], indent_dxa=120) -> None:
    table.alignment, table.autofit = WD_TABLE_ALIGNMENT.LEFT, False
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.find(qn("w:tblW"))
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(indent_dxa))
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for item in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(item))
        grid.append(node)


def add_code_block(doc: Document, code: str, language: str = "code") -> None:
    normalized = language.strip().lower() or "code"
    labels = {
        "python": ("PYTHON 3  ·  RUNNABLE", BLUE, "F4F7FA"),
        "py": ("PYTHON 3  ·  RUNNABLE", BLUE, "F4F7FA"),
        "output": ("EXECUTED OUTPUT  ·  BUILD VERIFIED", "1C8C8C", "F3F8F8"),
        "bash": ("TERMINAL  ·  REPRODUCIBLE COMMAND", INK, "F4F6F8"),
        "shell": ("TERMINAL  ·  REPRODUCIBLE COMMAND", INK, "F4F6F8"),
        "json": ("JSON  ·  STRUCTURED CONTRACT", DARK_BLUE, "F4F7FA"),
        "yaml": ("YAML  ·  CONFIGURATION", DARK_BLUE, "F4F7FA"),
        "text": ("TEXT / RESULT", MUTED, "F7F8FA"),
    }
    label, header_fill, body_fill = labels.get(
        normalized, (f"{normalized.upper()}  ·  CODE", DARK_BLUE, LIGHT_FILL)
    )
    table = doc.add_table(rows=2, cols=1)
    table_geometry(table, [PAGE_WIDTH_DXA])
    set_table_borders(table, color=header_fill, size=5)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_properties.append(OxmlElement("w:cantSplit"))
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)
    header = table.cell(0, 0)
    set_cell_width(header, PAGE_WIDTH_DXA)
    set_cell_margins(header, 50, 140, 45, 140)
    set_cell_shading(header, header_fill)
    header_paragraph = header.paragraphs[0]
    header_paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(
        header_paragraph.add_run(label),
        name="Calibri",
        size=8.0,
        color="FFFFFF",
        bold=True,
    )
    body = table.cell(1, 0)
    set_cell_width(body, PAGE_WIDTH_DXA)
    set_cell_margins(body, 110, 150, 110, 150)
    set_cell_shading(body, body_fill)
    paragraph = body.paragraphs[0]
    paragraph.paragraph_format.space_after, paragraph.paragraph_format.line_spacing = Pt(0), 1.0
    code_size = 8.35 if normalized not in {"output", "text"} else 8.2
    set_run_font(paragraph.add_run(code.rstrip()), name="Consolas", size=code_size, color="243746")


def add_table(doc: Document, rows: list[list[str]]) -> None:
    columns = max(len(row) for row in rows)
    rows = [row + [""] * (columns - len(row)) for row in rows]
    if columns == 2:
        widths = [2700, 6660]
    else:
        widths = [PAGE_WIDTH_DXA // columns] * columns
        widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=len(rows), cols=columns)
    set_table_borders(table)
    table_geometry(table, widths)
    for row_index, (word_row, values) in enumerate(zip(table.rows, rows, strict=False)):
        tr_pr = word_row._tr.get_or_add_trPr()
        if row_index == 0:
            tr_pr.append(OxmlElement("w:cantSplit"))
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for column, (cell, value) in enumerate(zip(word_row.cells, values, strict=False)):
            set_cell_width(cell, widths[column])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after, paragraph.paragraph_format.line_spacing = (
                Pt(0),
                1.0,
            )
            add_inline(paragraph, value.strip(), size=9.0 if columns >= 4 else 9.5)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_width(cell, PAGE_WIDTH_DXA)
    set_cell_margins(cell, 100, 180, 100, 180)
    set_cell_shading(cell, CALLOUT_FILL)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    add_inline(cell.paragraphs[0], text)


def add_figure(doc: Document, source: str, caption: str) -> None:
    path = ROOT / source
    if not path.exists():
        raise FileNotFoundError(f"Figure does not exist: {path}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    shape = paragraph.add_run().add_picture(str(path), width=Inches(6.15))
    set_picture_alt_text(shape, caption)
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.space_after = Pt(8)
    set_run_font(label.add_run(caption), size=9.5, color=MUTED, italic=True)


def add_equation(doc: Document, value: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True
    if EQUATION_CATALOG is None:
        raise RuntimeError("The native Word equation catalogue has not been initialised")
    EQUATION_CATALOG.append(paragraph, value)


def markdown_blocks(text: str):
    lines, i, paragraph = text.splitlines(), 0, []

    def flush():
        nonlocal paragraph
        if paragraph:
            value = " ".join(item.strip() for item in paragraph).strip()
            paragraph = []
            return "paragraph", value
        return None

    while i < len(lines):
        line = lines[i]
        if line.strip() == r"\[":
            item = flush()
            if item:
                yield item
            i += 1
            equation_lines = []
            while i < len(lines) and lines[i].strip() != r"\]":
                equation_lines.append(lines[i])
                i += 1
            yield "equation", "\n".join(equation_lines)
        elif line.startswith("```"):
            item = flush()
            if item:
                yield item
            language = line[3:].strip().lower() or "code"
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            yield "code", (language, "\n".join(code_lines))
        elif re.match(r"^!\[[^]]*\]\([^)]+\)$", line.strip()):
            item = flush()
            if item:
                yield item
            match = re.match(r"^!\[([^]]*)\]\(([^)]+)\)$", line.strip())
            assert match is not None
            yield "figure", (match.group(2), match.group(1))
        elif line.startswith("# "):
            item = flush()
            if item:
                yield item
            yield "h1", line[2:].strip()
        elif line.startswith("## "):
            item = flush()
            if item:
                yield item
            yield "h2", line[3:].strip()
        elif line.startswith("### "):
            item = flush()
            if item:
                yield item
            yield "h3", line[4:].strip()
        elif line.startswith("- "):
            item = flush()
            if item:
                yield item
            yield "bullet", line[2:].strip()
        elif re.match(r"^\d+\. ", line):
            item = flush()
            if item:
                yield item
            yield "number", re.sub(r"^\d+\. ", "", line).strip()
        elif line.startswith("> "):
            item = flush()
            if item:
                yield item
            yield "callout", line[2:].strip()
        elif (
            line.startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\|[\s:|-]+\|$", lines[i + 1])
        ):
            item = flush()
            if item:
                yield item
            rows = [[c.strip() for c in line.strip().strip("|").split("|")]]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            yield "table", rows
            continue
        elif not line.strip():
            item = flush()
            if item:
                yield item
        else:
            paragraph.append(line)
        i += 1
    item = flush()
    if item:
        yield item


def normalize_equation(value: str) -> str:
    """Return the stable key used for source TeX and generated Word mathematics."""

    return " ".join(line.strip() for line in value.splitlines()).strip()


def prepare_equation_for_word(value: str) -> str:
    """Typeset domain acronyms as named quantities rather than letter products."""

    protected: list[str] = []

    def protect(match: re.Match[str]) -> str:
        protected.append(match.group(0))
        return f"@@CRBPROTECTED{len(protected) - 1}@@"

    prepared = TEXT_MATH_PATTERN.sub(protect, normalize_equation(value))
    prepared = ROMAN_MATH_PATTERN.sub(lambda match: rf"\text{{{match.group(1)}}}", prepared)
    for index, original in enumerate(protected):
        prepared = prepared.replace(f"@@CRBPROTECTED{index}@@", original)
    return prepared


def _inline_equations(text: str):
    for match in INLINE.finditer(text):
        token = match.group(0)
        if token.startswith("$"):
            yield normalize_equation(token[1:-1])


def collect_source_equations() -> list[str]:
    """Collect equations from rendered prose while deliberately ignoring code fences."""

    equations: set[str] = set()
    for directory in (MANUSCRIPT, GUIDED_LABS):
        for path in sorted(directory.glob("*.md")):
            for kind, value in markdown_blocks(path.read_text(encoding="utf-8")):
                if kind == "equation":
                    equations.add(normalize_equation(value))
                elif kind == "table":
                    for row in value:
                        for cell in row:
                            equations.update(_inline_equations(cell))
                elif kind in {
                    "paragraph",
                    "bullet",
                    "number",
                    "callout",
                    "h1",
                    "h2",
                    "h3",
                }:
                    equations.update(_inline_equations(value))
    if not equations:
        raise RuntimeError("No source equations were found")
    return sorted(equations)


class PandocEquationCatalog:
    """Compile TeX once and reuse native, editable Office Math objects.

    Pandoc translates each source expression to OMML, Word's native equation
    representation. This preserves fractions, radicals, matrices, n-ary
    operators and limits; it also avoids the former lossy Sigma substitution.
    """

    def __init__(self, equations: list[str]) -> None:
        self._elements: dict[str, etree._Element] = {}
        markdown = "\n\n".join(
            f"EQ{index:04d}\n\n$$\n{prepare_equation_for_word(equation)}\n$$"
            for index, equation in enumerate(equations)
        )
        with tempfile.TemporaryDirectory(prefix="credit-risk-equations-") as directory:
            directory_path = Path(directory)
            source = directory_path / "equations.md"
            output = directory_path / "equations.docx"
            source.write_text(markdown, encoding="utf-8")
            try:
                subprocess.run(
                    [
                        "pandoc",
                        str(source),
                        "--from=markdown+tex_math_dollars",
                        "--to=docx",
                        f"--output={output}",
                    ],
                    check=True,
                    capture_output=True,
                    text=True,
                )
            except FileNotFoundError as exc:
                raise RuntimeError(
                    "Pandoc is required to build native Word equations; install pandoc first"
                ) from exc
            except subprocess.CalledProcessError as exc:
                raise RuntimeError(f"Pandoc equation conversion failed: {exc.stderr}") from exc
            with zipfile.ZipFile(output) as archive:
                root = etree.fromstring(archive.read("word/document.xml"))

        namespaces = {"m": MATH_NAMESPACE, "w": WORD_NAMESPACE}
        generated: dict[int, etree._Element] = {}
        current_index: int | None = None
        for paragraph in root.xpath("//w:body/w:p", namespaces=namespaces):
            visible_text = "".join(paragraph.xpath(".//w:t/text()", namespaces=namespaces))
            marker = re.fullmatch(r"EQ(\d{4})", visible_text.strip())
            if marker:
                current_index = int(marker.group(1))
                continue
            math_objects = paragraph.xpath(".//m:oMath", namespaces=namespaces)
            if math_objects and current_index is not None:
                if len(math_objects) != 1:
                    raise RuntimeError(
                        f"Expected one Word math object for equation {current_index}, "
                        f"found {len(math_objects)}"
                    )
                generated[current_index] = math_objects[0]
                current_index = None
        if len(generated) != len(equations):
            missing = [
                equation for index, equation in enumerate(equations) if index not in generated
            ]
            raise RuntimeError(
                f"Equation conversion count mismatch: {len(equations)} source expressions, "
                f"{len(generated)} Word equations. Missing: {missing[:3]}"
            )
        for index, equation in enumerate(equations):
            element = generated[index]
            self._audit_operator_structure(equation, element)
            self._elements[equation] = copy.deepcopy(element)

    @staticmethod
    def _audit_operator_structure(equation: str, element: etree._Element) -> None:
        xml = etree.tostring(element, encoding="unicode")
        requirements = {
            r"\sum": ("∑", "summation"),
            r"\prod": ("∏", "product"),
            r"\frac": ("<m:f>", "fraction"),
            r"\sqrt": ("<m:rad>", "radical"),
        }
        for command, (marker, label) in requirements.items():
            if command in equation and marker not in xml:
                raise RuntimeError(f"Native Word {label} structure missing for: {equation}")

    def append(self, paragraph, equation: str) -> None:
        key = normalize_equation(equation)
        try:
            element = self._elements[key]
        except KeyError as exc:
            raise KeyError(f"Equation was not precompiled: {key}") from exc
        paragraph._p.append(copy.deepcopy(element))


def _chapter_lab_blocks(chapter_number: int):
    path = GUIDED_LABS / f"chapter_{chapter_number:02d}.md"
    if not path.exists():
        raise FileNotFoundError(f"Missing guided laboratory: {path}")
    yield from markdown_blocks(path.read_text(encoding="utf-8"))


def add_manuscript(doc: Document, bullet_num: int, decimal_num: int) -> None:
    structure = json.loads(STRUCTURE.read_text(encoding="utf-8"))
    roman = ("I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII")
    parts = {
        part["chapters"][0]["number"]: (f"Part {roman[index]}", part["title"])
        for index, part in enumerate(structure["parts"])
    }
    previous_kind = None
    current_decimal_num = decimal_num
    current_chapter: int | None = None
    render_chapter: int | None = None
    section_number = 0
    subsection_number = 0
    bookmark_id = 1
    rendered_part_count = 0
    for path in sorted(MANUSCRIPT.glob("*.md")):
        content = path.read_text(encoding="utf-8")
        if path.name == "00_front_matter.md":
            content = content[content.index("## Preface") :]
        blocks = list(markdown_blocks(content))
        expanded_blocks = []
        for kind, value in blocks:
            if kind == "h1":
                match = re.match(r"Chapter (\d+)\b", value)
                if match and current_chapter is not None:
                    expanded_blocks.extend(_chapter_lab_blocks(current_chapter))
                current_chapter = int(match.group(1)) if match else None
            expanded_blocks.append((kind, value))
        if current_chapter is not None and path.name.startswith("12_"):
            expanded_blocks.extend(_chapter_lab_blocks(current_chapter))
            current_chapter = None
        for kind, value in expanded_blocks:
            if kind == "h1":
                match = re.match(r"Chapter (\d+)\b", value)
                render_chapter = int(match.group(1)) if match else None
                section_number = 0
                subsection_number = 0
                if match and int(match.group(1)) in parts:
                    add_part_banner(
                        doc,
                        *parts[int(match.group(1))],
                        first_part=rendered_part_count == 0,
                    )
                    rendered_part_count += 1
                paragraph = doc.add_paragraph(value, style="Heading 1")
                if match:
                    add_bookmark(paragraph, f"chapter_{int(match.group(1)):03d}", bookmark_id)
                    bookmark_id += 1
                elif value in {
                    "Appendices",
                    "Practice Casebook — Seventy-Two Worked Assignments",
                    "Technical Workbook — End-to-End Python Patterns",
                    "Numerical Examples — Calculation, Interpretation, and Audit",
                    "Credit Risk Policy Playbook",
                    "Review and Viva Questions with Instructor Notes",
                    "Technical and Governance Glossary",
                }:
                    anchor = re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")
                    add_bookmark(paragraph, f"section_{anchor}", bookmark_id)
                    bookmark_id += 1
                paragraph.paragraph_format.page_break_before = False
            elif kind == "h2":
                section_number += 1
                subsection_number = 0
                heading_text = re.sub(r"^\d+\.\s+", "", value)
                label = (
                    f"{render_chapter}.{section_number} {heading_text}"
                    if render_chapter
                    else heading_text
                )
                doc.add_paragraph(label, style="Heading 2")
            elif kind == "h3":
                subsection_number += 1
                heading_text = re.sub(r"^\d+\.\s+", "", value)
                label = (
                    f"{render_chapter}.{section_number}.{subsection_number} {heading_text}"
                    if render_chapter
                    else heading_text
                )
                doc.add_paragraph(label, style="Heading 3")
            elif kind == "paragraph":
                paragraph = doc.add_paragraph()
                if value.startswith("[R"):
                    (
                        paragraph.paragraph_format.space_after,
                        paragraph.paragraph_format.line_spacing,
                    ) = Pt(5), 1.15
                    add_inline(paragraph, value, size=9.5)
                else:
                    add_inline(paragraph, value)
            elif kind in {"bullet", "number"}:
                if kind == "number" and previous_kind != "number":
                    current_decimal_num = add_custom_numbering(doc, bullet=False)
                paragraph = doc.add_paragraph()
                apply_numbering(paragraph, bullet_num if kind == "bullet" else current_decimal_num)
                add_inline(paragraph, value)
            elif kind == "code":
                language, code = value
                add_code_block(doc, code, language)
            elif kind == "table":
                add_table(doc, value)
            elif kind == "callout":
                add_callout(doc, value)
            elif kind == "equation":
                add_equation(doc, value)
            elif kind == "figure":
                add_figure(doc, *value)
            previous_kind = kind


def audit_document(doc: Document) -> None:
    section = doc.sections[0]
    assert section.page_width == Inches(8.5) and section.page_height == Inches(11)
    assert section.left_margin == Inches(1)
    assert doc.styles["Normal"].font.name == "Calibri" and doc.styles["Normal"].font.size == Pt(11)
    assert doc.styles["Heading 1"].font.size == Pt(16)
    assert len(doc.tables) >= 10
    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert any(text.startswith("Chapter 72") for text in headings)
    assert "Appendices" in headings and any(
        text.startswith("Practice Casebook") for text in headings
    )


def build(output: Path) -> None:
    global EQUATION_CATALOG
    EQUATION_CATALOG = PandocEquationCatalog(collect_source_equations())
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    bullet_num, decimal_num = (
        add_custom_numbering(doc, bullet=True),
        add_custom_numbering(doc, bullet=False),
    )
    doc.core_properties.title = "Intelligent Credit Risk Modeling with Python"
    doc.core_properties.subject = (
        "From Data Quality and Scorecards to IFRS 9, Basel IRB, Deployment, and Governed Agentic AI"
    )
    doc.core_properties.author = "Dr. Ferdinantos Kottas"
    doc.core_properties.keywords = "credit risk, scorecard, IFRS 9, IRB, agentic AI, Python"
    add_cover(doc)
    add_toc(doc)
    add_manuscript(doc, bullet_num, decimal_num)
    update = doc.settings._element.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        doc.settings._element.append(update)
    update.set(qn("w:val"), "true")
    audit_document(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        type=Path,
        default=ROOT
        / "artifacts"
        / "Intelligent_Credit_Risk_Modeling_with_Python_Analytical_Review.docx",
    )
    build(parser.parse_args().output.resolve())


if __name__ == "__main__":
    main()
